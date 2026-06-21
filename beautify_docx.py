"""美化 + 移除 NotebookLM / Gemini 关键字，一步到位。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

SRC = r"F:\AeroTutor_设计说明书.docx"
DST = r"F:\AeroTutor_设计说明书_美化版.docx"

doc = Document(SRC)

# ── 替换规则：去掉 NotebookLM / Gemini ──────────────────
# 先做精确短语替换，再做通用关键词清理
PHRASE_REPLACEMENTS = [
    # 从最具体到最通用排列
    ("是一款 NotebookLM 风格的多知识库自适应学习 Agent", "是一款多知识库自适应学习 Agent"),
    ("是一款 NotebookLM 风格的", "是一款"),
    ("是一款 NotebookLM 风格 ", "是一款"),
    ("NotebookLM 风格的多知识库自适应学习 Agent", "多知识库自适应学习 Agent"),
    ("NotebookLM 风格的", ""),
    ("（NotebookLM 风格）", ""),
    ("NotebookLM / Gemini 风格的", ""),
    ("NotebookLM / Gemini 风格", ""),
    ("NotebookLM 风格 ", ""),
    ("NotebookLM 风格", ""),
    ("Gemini 风格气泡", "对话气泡"),
    ("Gemini 浅色主题", "浅色主题"),
    ("Gemini 风格", ""),
    ("Gemini ", ""),
]

def strip_keywords(text):
    """移除 NotebookLM / Gemini 关键字并清理残留空格标点。"""
    for old, new in PHRASE_REPLACEMENTS:
        text = text.replace(old, new)
    # 极端情况：仍有独立出现的 NotebookLM 或 Gemini
    text = text.replace("NotebookLM", "")
    text = text.replace("Gemini", "")
    # 清理残留：连续空格、多余斜杠、"的的"、句首多余标点
    text = re.sub(r' {2,}', ' ', text)
    text = re.sub(r'\s*/\s*', '', text)       # 去掉 " / " 残留
    text = re.sub(r'的的', '的', text)         # 修复 "的的"
    text = re.sub(r'，的', '，', text)         # 修复 "，的"
    text = re.sub(r'\s+的\s+', '', text)       # 去掉孤立的 "的"（前后有空格）
    text = re.sub(r'。的', '。', text)         # 修复 "。的"
    text = re.sub(r'、的', '、', text)         # 修复 "、的"
    return text.strip()

# ── 先做文本替换 ───────────────────────────────────────
for p in doc.paragraphs:
    if not p.runs:
        continue
    full = "".join(r.text for r in p.runs)
    new_text = strip_keywords(full)
    if new_text != full:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""

# ── 全局样式 ───────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1F, 0x29, 0x33)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(4)
pf.space_after = Pt(4)

for level, (size, color, bold) in {
    "Heading 1": (Pt(18), RGBColor(0x1A, 0x3A, 0x5C), True),
    "Heading 2": (Pt(15), RGBColor(0x2E, 0x86, 0xC1), True),
    "Heading 3": (Pt(13), RGBColor(0x2E, 0x86, 0xC1), True),
}.items():
    s = doc.styles[level]
    s.font.name = "微软雅黑"
    s.font.size = size
    s.font.bold = bold
    s.font.color.rgb = color
    s.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    pf = s.paragraph_format
    pf.line_spacing = 1.2
    if level == "Heading 1":
        pf.space_before = Pt(18)
        pf.space_after = Pt(10)
        pPr = s.element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "8",
            qn("w:space"): "4", qn("w:color"): "2E86C1",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == "Heading 2":
        pf.space_before = Pt(14)
        pf.space_after = Pt(6)
    else:
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)

for name in ("List Bullet", "List Bullet 2", "List Number"):
    try:
        s = doc.styles[name]
    except KeyError:
        continue
    s.font.name = "微软雅黑"
    s.font.size = Pt(11)
    s.font.color.rgb = RGBColor(0x1F, 0x29, 0x33)
    s.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    pf = s.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)

# ── 逐段微调 ───────────────────────────────────────────
for idx, p in enumerate(doc.paragraphs):
    pStyle = p.style.name
    text = p.text.strip()

    for r in p.runs:
        if r.font.name is None:
            r.font.name = "微软雅黑"
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = rPr.makeelement(qn("w:rFonts"), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "微软雅黑")
        if pStyle in ("Normal", "List Bullet", "List Bullet 2", "List Number"):
            if r.font.size and Pt(10) < r.font.size < Pt(14):
                r.font.size = Pt(11)

    # 封面大标题
    if pStyle == "Normal" and idx < 12 and text == "沈航智学 AeroTutor":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(60)
        p.paragraph_format.space_after = Pt(8)
        for r in p.runs:
            r.font.size = Pt(30)
            r.font.bold = True

    # 封面副标题
    if pStyle == "Normal" and len(text) < 60 and "自适应学习" in text:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        for r in p.runs:
            r.font.size = Pt(16)

    # "设计说明书"
    if pStyle == "Normal" and text == "设计说明书":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(13)

    # 学校 + 大赛
    if pStyle == "Normal" and len(text) < 80 and "筑梦空天" in text:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(36)
        for r in p.runs:
            r.font.size = Pt(12)

    # 尾页标题
    if idx >= len(doc.paragraphs) - 3 and text == "沈航智学 AeroTutor":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(40)
        for r in p.runs:
            r.font.size = Pt(20)
            r.font.bold = True

    # 尾页学校
    if idx >= len(doc.paragraphs) - 2 and "沈阳航空航天大学" in text:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(12)

    # 目录项
    if pStyle == "Normal" and text and text[0] in "一二三四五六七八九十" and "、" in text[:3]:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)

    # 创新点标题加粗
    if pStyle == "Normal" and "创新点" in text and len(text) < 60:
        for r in p.runs:
            r.font.bold = True

    # 空段落压缩
    if text == "" and pStyle == "Normal":
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

# ── 页边距 ─────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin = Pt(54)
    sec.bottom_margin = Pt(54)
    sec.left_margin = Pt(72)
    sec.right_margin = Pt(72)

doc.save(DST)
print(f"Done: {DST}")
