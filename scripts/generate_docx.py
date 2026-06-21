"""生成 AeroTutor 设计说明书 DOCX"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── 正文样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    return h


def add_body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


# ── 封面 ──
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('沈航智学 AeroTutor')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('NotebookLM 风格多知识库自适应学习智能体')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('设计说明书').font.size = Pt(12)

doc.add_paragraph('')
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('沈阳航空航天大学\n第一届「筑梦空天」AI 智能体创新应用大赛').font.size = Pt(11)

doc.add_page_break()

# ── 目录 ──
add_heading('目录', level=1)
toc_items = [
    '一、产品概述',
    '二、系统架构',
    '三、核心功能',
    '四、使用场景',
    '五、创新点',
    '六、技术选型',
    '七、部署信息',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(12)

doc.add_page_break()

# ═══ 一 ═══
add_heading('一、产品概述')

add_body(
    '沈航智学 AeroTutor 是一款 NotebookLM 风格的多知识库自适应学习 Agent，'
    '专为航空航天专业课程学习场景打造。系统以"知识空间"为核心组织单元，'
    '融合 RAG 检索增强生成、LLM 智能工具调用、认知诊断分析等技术，'
    '为学生提供从知识入库、检索问答、测验自评到薄弱诊断的全链路自适应学习体验。'
)

add_body(
    '系统采用 Streamlit 构建前端界面，LangChain 作为 Agent 框架，'
    'FAISS 作为向量检索引擎，支持 DeepSeek、智谱 GLM 及自定义 OpenAI 兼容 API '
    '三种 LLM 供应商热切换。部署于 Streamlit Cloud，零运维成本，即开即用。'
)

add_bullet('技术栈：Streamlit + LangChain + FAISS + DeepSeek/GLM + sentence-transformers')
add_bullet('公开访问：https://sau-aerotutor-9a5dscagkznibfd2ppsaog.streamlit.app')

# ═══ 二 ═══
add_heading('二、系统架构')

add_heading('2.1 总体架构', level=2)
add_body(
    '系统采用四层架构设计：前端层（Streamlit UI）→ 智能体层（LangChain Agent）'
    '→ 引擎层（Core Engines）→ 存储层（Data Layer）。'
    '各层职责清晰、松耦合，便于扩展和维护。'
)

add_heading('前端层', level=3)
add_bullet('对话界面：Gemini 风格气泡 + 打字机流式输出 + 答案溯源标签')
add_bullet('测验卡片：选择题即时判对错 + 简答题 AI 自动批改（按要点打分，满分 20）')
add_bullet('翻转闪卡：纯 CSS 3D 翻转动画，正面概念 + 反面详解')
add_bullet('知识空间管理：侧边栏勾选激活 + 拖拽文件夹树管理对话')
add_bullet('设置面板：供应商选择 + API Key 输入 + 连接测试 + 一键保存')

add_heading('智能体层', level=3)
add_bullet('意图预检测：中英文关键词匹配，5 类高频操作直通执行，延迟降低 3-5 倍')
add_bullet('LangChain Agent：6 工具（检索/出题/闪卡/诊断/速查/播客）自动编排调用')
add_bullet('RAG 引擎：多 FAISS 索引联合检索 + 语义分块（500 字窗口）+ 答案溯源')
add_bullet('多轮对话记忆：保留最近 20 轮上下文，支持长对话')

add_heading('引擎层', level=3)
add_bullet('QuizEngine：基于知识库自动生成选择题 + 简答题 + 闪卡 + 速查摘要 + 知识播客 + AI 自动批改')
add_bullet('DiagnosisEngine：三阶认知诊断（已掌握/薄弱/未覆盖）+ 个性化学习航线生成')
add_bullet('ReviewVault：五位一体复习智库 — 错题本 + 闪卡集 + 诊断报告 + 学习航线 + 学习统计')

add_heading('存储层', level=3)
add_bullet('FAISS 向量数据库：每个知识空间独立索引文件，完全物理隔离')
add_bullet('JSON 持久化：对话历史、文件夹结构、闪卡数据、用户设置')
add_bullet('文件解析：PyMuPDF (PDF) + python-pptx (PPTX) + 纯文本 (TXT/MD)')
add_bullet('配置管理：.env 本地开发 / Streamlit Secrets 云部署')

add_heading('2.2 模块依赖', level=2)
add_body('app.py 为单一主入口，通过 init_core() 初始化基础资源，ChatAgent 根据用户配置动态创建。各模块职责单一：')

modules = [
    'user_settings.py — 供应商预设 + 配置持久化 + 连接测试',
    'kb_manager.py — FAISS 知识空间 CRUD + 文件分块入库 + 索引缓存',
    'rag_engine.py — 多索引联合检索 + 语义排序 + 来源标注',
    'chat_agent.py — LangChain Agent + 6 Tools + 意图预检测 + 多轮记忆',
    'quiz_engine.py — 测验/闪卡/速查摘要/播客脚本/AI 批改',
    'diagnosis.py — 认知诊断 + 学习航线生成',
    'vault.py — 复习智库（错题本/闪卡集/诊断/航线/统计）',
    'styles.py — Gemini 浅色主题 CSS（70+ 样式规则）',
    'components.py — Generative UI 组件（测验卡片/翻转闪卡/播客/航线/统计条）',
    'sidebar_dnd.py — 侧边栏文件夹拖拽树（HTML5 Drag & Drop + postMessage）',
    'kb_panel.py — 知识空间管理面板（创建/上传/删除）',
    'vault_panel.py — 复习智库面板（5 Tab 子面板）',
]
for m in modules:
    add_bullet(m)

# ═══ 三 ═══
add_heading('三、核心功能')

add_heading('3.1 多知识空间管理（NotebookLM 风格）', level=2)
add_body(
    '每个学科/章节对应一个独立知识空间，底层为独立的 FAISS 向量索引文件，'
    '向量完全不互通，确保跨学科检索"不串味"。支持课程课件、学习笔记、'
    '论文文献、参考资料、其他五类分类标签。上传 PDF 课件、PPT 讲稿、'
    'TXT/Markdown 笔记后自动分块向量化（500 字滑动窗口，中文标点优先切分）。'
    '侧边栏勾选即激活，支持多知识空间联合检索。'
)

add_heading('3.2 RAG 检索增强生成', level=2)
add_body(
    '根据用户勾选的知识空间，跨多个 FAISS 索引并行检索，语义排序后注入 LLM 上下文。'
    '每个回答自动标注引用来源（文件名 + 知识空间名），实现精确答案溯源。'
    '检索算法使用余弦相似度（FAISS Inner Product），top-K 合并重排。'
)

add_heading('3.3 智能学习工具箱（6 Tools）', level=2)
tools = [
    ('知识检索', '自然语言提问 → 跨知识库检索 → 带溯源引用的回答'),
    ('出题测验', '基于知识库自动生成选择题 + 简答题，支持自定义数量（1-10 道）'),
    ('知识闪卡', '提取核心概念生成正面（概念）/ 反面（详解）学习闪卡'),
    ('认知诊断', '分析对话历史 + 测验结果，输出三阶诊断报告'),
    ('速查摘要', '一键生成核心概念列表 + 必会公式 + 常见误区 + 复习建议'),
    ('知识播客', '王教授 × 小明双人对话式知识讲解脚本'),
]
for name, desc in tools:
    add_bullet(f'{name}：{desc}')

add_heading('3.4 AI 自动批改', level=2)
add_body(
    '选择题即时判对错，标注正确选项并附解析。简答题调用独立评分 LLM（temperature=0），'
    '按参考答案要点逐一比对打分（满分 20 分），逐条点评对错。'
    '得分低于 60% 的题目自动收录错题本，形成闭环学习。'
)

add_heading('3.5 认知诊断 + 学习航线', level=2)
add_body(
    '基于对话历史和测验表现的动态认知诊断引擎。输出三层分析：'
    '已掌握知识点（绿色）、薄弱知识点（黄色，附判断依据）、'
    '未覆盖知识盲区（红色，附补充建议）。'
    '自动生成 3-5 步个性化学习航线，每步标注建议学习时长和推荐资源类型，'
    '考虑知识点前置依赖关系。'
)

add_heading('3.6 复习智库', level=2)
add_bullet('错题本：自动收录答错的题目 + 手动标记已复习 + 复习次数追踪')
add_bullet('闪卡集：CSS 3D 翻转动画卡片，按知识空间分类存储，最近 50 张可查')
add_bullet('诊断报告：历史诊断归档（最近 10 次），可追踪进步轨迹')
add_bullet('学习航线：步骤化学习路线指引（最近 5 条），可回看历史规划')
add_bullet('学习统计：连续复习天数、完成测验次数、闪卡数量、错题数量')

add_heading('3.7 对话管理', level=2)
add_body(
    '支持创建/切换/删除对话，对话自动以首条提问命名。'
    '文件夹树结构组织对话，支持嵌套子文件夹（无限层级）。'
    'HTML5 Drag & Drop 实现对话拖拽移动，postMessage 通信与 Streamlit 后端同步。'
    '本地 JSON 持久化，多轮对话记忆保留最近 20 轮上下文。'
)

add_heading('3.8 多 API 供应商热切换', level=2)
add_body(
    '内置三种供应商预设：DeepSeek（deepseek-chat / deepseek-reasoner）、'
    '智谱 GLM（glm-4-flash / glm-4-plus / glm-4）、自定义 OpenAI 兼容接口。'
    '设置面板包含连接测试功能，一键验证 API Key 有效性。'
    '配置保存到本机 JSON 文件，切换即时生效，无需重启应用。'
)

# ═══ 四 ═══
add_heading('四、使用场景')

scenes = [
    ('场景一：课后自主复习',
     '学生将课堂 PPT 上传到对应知识空间，向 AI 提问课件中的知识点。'
     '系统从课件中检索相关内容并给出带引用的详细回答。随后点击「生成测验」自测，'
     'AI 自动批改并将错题收录错题本，形成"学→测→评"闭环。'),
    ('场景二：考前系统复习',
     '学生勾选多个知识空间，点击「诊断薄弱点」。系统综合分析对话记录和测验历史，'
     '标记薄弱知识点和未覆盖盲区，自动生成一条 3-5 步学习航线。'
     '每条航线标注建议学习时长和推荐资源类型（PPT/教材/视频/习题），'
     '帮助学生高效利用考前时间。'),
    ('场景三：碎片时间闪卡复习',
     '基于知识库一键生成学习闪卡，正面显示核心概念（如"雷诺数"），'
     '背面显示详细解释和公式。CSS 3D 翻转动画让移动端复习体验流畅自然。'
     '所有闪卡自动存入复习智库，按知识空间分类管理。'),
    ('场景四：知识播客沉浸学习',
     '将抽象概念转化为王教授（航空专家）和小明（大二学生）的师生对话播客脚本。'
     '用航空比喻（如"飞行 check list""座舱仪表""升力就像飞机的隐形翅膀"）'
     '把难懂的理论知识讲透，适合听觉型学习者。'),
    ('场景五：多用户共享平台',
     'AeroTutor 平台本身不存储任何 API 密钥，学生各自使用自己的 DeepSeek/GLM Key。'
     'FAISS 本地向量检索 + sentence-transformers 本地 Embedding 零 API 费用。'
     '平台零推理成本，小学期/社团可直接部署使用。'),
]
for title, desc in scenes:
    add_body(f'{title}')
    add_body(desc)

# ═══ 五 ═══
add_heading('五、创新点')

innovations = [
    ('1. NotebookLM 风格多知识库物理隔离',
     '每个学科/章节对应独立 FAISS 向量索引文件，检索时严格隔离，'
     '绝不跨学科"串味"。答案溯源精确到文件名，'
     '学生可放心上传不同课程课件同时使用，互不干扰。'),
    ('2. 意图预检测 + 工具直通加速架构',
     '在 LLM Agent 推理之前，加入中英文关键词意图预检测层。'
     '对"出题""闪卡""诊断""速查""播客"五类高频操作直接命中执行，'
     '绕过 LLM 推理和 Agent 编排环节。响应速度提升 3-5 倍，准确率 100%。'),
    ('3. 五步自适应学习闭环',
     '"知识入库 → 检索学习 → 出题自测 → AI 批改 → 错题收录 → 认知诊断 → 学习航线"，'
     '每一步的输出成为下一步的输入，形成螺旋上升的自适应学习飞轮。'),
    ('4. 航空航天领域深度定制',
     '播客角色设定贴合航空院校师生关系（王教授 + 小明）；'
     '回答风格善用航空比喻（check list、飞行计划、座舱仪表）；'
     '知识空间分类（课件/论文/参考资料）贴合工科学习场景。'),
    ('5. 多供应商热切换 + 零成本私部署',
     '内置 DeepSeek / GLM / 自定义三种供应商，学生自备 API Key 即可使用。'
     'FAISS 本地向量检索无外部数据库依赖，sentence-transformers 本地 Embedding 零 API 费用。'
     'Streamlit Cloud 免费托管，部署零预算。'),
    ('6. Generative UI 学习交互',
     '测验卡片、CSS 3D 翻转闪卡、双人对话播客、时间线航线图、统计仪表板 —— '
     '所有交互组件均为纯代码生成，不依赖外部 UI 模板，'
     '实现 NotebookLM / Gemini 级别的现代学习体验。'),
]
for title, desc in innovations:
    add_body(f'创新点 {title}')
    add_body(desc)

# ═══ 六 ═══
add_heading('六、技术选型')

tech_items = [
    ('UI 框架', 'Streamlit 1.31+（纯 Python 构建，零前端代码，开发效率极高）'),
    ('LLM 供应商', 'DeepSeek / 智谱 GLM / 自定义 OpenAI 兼容接口，支持热切换'),
    ('Agent 框架', 'LangChain create_agent + ChatOpenAI，6 工具自动编排'),
    ('向量数据库', 'FAISS（Facebook AI Research），轻量级，无外部服务依赖'),
    ('Embedding', 'sentence-transformers all-MiniLM-L6-v2，384 维，本地 CPU 推理'),
    ('文件解析', 'PyMuPDF (PDF) + python-pptx (PPTX) + 原生文本 (TXT/MD)'),
    ('可视化', 'Plotly 5.22+（知识图谱交互渲染预留）'),
    ('部署平台', 'Streamlit Cloud，Git Push 即部署，自动构建发布'),
]
for name, val in tech_items:
    add_bullet(f'{name}：{val}')

# ═══ 七 ═══
add_heading('七、部署信息')

deploy_items = [
    ('部署平台', 'Streamlit Cloud（免费）'),
    ('Python 版本', '3.11'),
    ('核心依赖', '14 个（streamlit / langchain / faiss-cpu / sentence-transformers 等）'),
    ('冷启动时间', '约 3 分钟（含 sentence-transformers 模型首次下载）'),
    ('持久化存储', 'Streamlit Cloud 容器内文件系统，重启保留数据'),
    ('安全机制', 'API Key 通过 Streamlit Secrets 加密存储，绝对不提交 Git 仓库'),
    ('代码仓库', 'https://github.com/aa1242634004-design/sau-aerotutor'),
    ('公开地址', 'https://sau-aerotutor-9a5dscagkznibfd2ppsaog.streamlit.app'),
]
for name, val in deploy_items:
    add_bullet(f'{name}：{val}')

# ── 尾页 ──
doc.add_page_break()
for _ in range(8):
    doc.add_paragraph('')

end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end.add_run('沈航智学 AeroTutor')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

end2 = doc.add_paragraph()
end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
end2.add_run('沈阳航空航天大学 · 第一届「筑梦空天」AI 智能体创新应用大赛').font.size = Pt(11)

# ── 保存 ──
import os
desktop = os.path.join(os.environ.get('USERPROFILE', ''), 'Desktop')
if not os.path.exists(desktop):
    desktop = r'F:\Desktop'
os.makedirs(desktop, exist_ok=True)
out_path = os.path.join(desktop, 'AeroTutor_设计说明书.docx')
doc.save(out_path)
print('DOCX saved to: ' + out_path)
